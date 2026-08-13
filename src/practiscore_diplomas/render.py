"""Render diploma-data YAML records into DOCX, ODT, or PDF documents."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any, Mapping
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

import yaml
from docx import Document
from docx.document import Document as DocumentObject
from docx.opc.exceptions import PackageNotFoundError
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_BREAK



class DiplomaRenderError(ValueError):
    """Raised when diploma data or a document template cannot be rendered."""


_PLACEHOLDER = re.compile(r"{{\s*([^{}]+?)\s*}}")
_TOP_LEVEL_FIELDS = {
    "first_line",
    "second_line",
    "third_line",
    "fourth_line",
    "place",
    "division",
    "category",
    "class",
    "metric_value",
}


def _records_from_data(raw: Any) -> list[Mapping[str, Any]]:
    if not isinstance(raw, dict) or not isinstance(raw.get("diplomas"), dict):
        raise DiplomaRenderError("Diploma data must contain a 'diplomas' mapping")
    records: list[Mapping[str, Any]] = []
    for series, values in raw["diplomas"].items():
        if not isinstance(values, list):
            raise DiplomaRenderError(f"diplomas.{series} must be a list")
        for index, record in enumerate(values):
            if not isinstance(record, dict):
                raise DiplomaRenderError(f"diplomas.{series}[{index}] must be a mapping")
            records.append(record)
    if not records:
        raise DiplomaRenderError("Diploma data contains no diploma records")
    return records


def _load_records(path: Path) -> list[Mapping[str, Any]]:
    try:
        with path.open("r", encoding="utf-8") as stream:
            raw = yaml.safe_load(stream)
    except (OSError, yaml.YAMLError) as exc:
        raise DiplomaRenderError(f"Could not read diploma data {path}: {exc}") from exc
    return _records_from_data(raw)


def _resolve(record: Mapping[str, Any], path: str, record_index: int) -> Any:
    parts = path.split(".")
    if parts[0] == "shooter":
        if len(parts) != 2 or not isinstance(record.get("shooter"), dict):
            raise DiplomaRenderError(f"Record {record_index}: invalid placeholder '{{{{ {path} }}}}'")
        value = record["shooter"].get(parts[1])
    elif len(parts) == 1 and parts[0] in _TOP_LEVEL_FIELDS:
        value = record.get(parts[0])
        if parts[0] in {"third_line", "fourth_line"} and parts[0] not in record:
            return ""
    else:
        raise DiplomaRenderError(f"Record {record_index}: unknown placeholder '{{{{ {path} }}}}'")
    if value is None:
        raise DiplomaRenderError(f"Record {record_index}: placeholder '{{{{ {path} }}}}' has no value")
    if isinstance(value, list):
        return ", ".join(str(item) for item in value)
    return value


def _render_text(text: str, record: Mapping[str, Any], record_index: int) -> str:
    if text.count("{{") != text.count("}}"):
        raise DiplomaRenderError(f"Record {record_index}: malformed placeholder syntax")

    def replace(match: re.Match[str]) -> str:
        path = match.group(1).strip()
        value = _resolve(record, path, record_index)
        return str(value)

    return _PLACEHOLDER.sub(replace, text)


def _replace_paragraph(paragraph: Paragraph, record: Mapping[str, Any], record_index: int) -> None:
    runs = list(paragraph.runs)
    if not runs:
        return
    combined = "".join(run.text or "" for run in runs)
    if "{{" not in combined and "}}" not in combined:
        return
    rendered = _render_text(combined, record, record_index)
    runs[0].text = rendered
    for run in runs[1:]:
        run.text = ""


def _replace_cell(cell: _Cell, record: Mapping[str, Any], record_index: int) -> None:
    for paragraph in cell.paragraphs:
        _replace_paragraph(paragraph, record, record_index)
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                _replace_cell(nested_cell, record, record_index)


def _render_document_page(document: DocumentObject, record: Mapping[str, Any], record_index: int) -> None:
    for paragraph in document.paragraphs:
        _replace_paragraph(paragraph, record, record_index)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_cell(cell, record, record_index)


def _body_blocks(document: DocumentObject) -> list[Any]:
    body = document._body._element
    return [deepcopy(child) for child in body if child.tag.endswith("}p") or child.tag.endswith("}tbl")]


def _insert_before_section_properties(body: Any, block: Any) -> None:
    section_properties = body.sectPr
    if section_properties is None:
        body.append(block)
    else:
        body.insert(body.index(section_properties), block)


def _render_records(records: list[Mapping[str, Any]], template_path: str | Path, output_path: str | Path) -> None:
    template = Path(template_path)
    if not template.is_file():
        raise DiplomaRenderError(f"Template does not exist: {template}")
    output = Path(output_path)
    try:
        result = Document(template)
    except (OSError, ValueError, KeyError, PackageNotFoundError) as exc:
        raise DiplomaRenderError(f"Could not open DOCX template {template}: {exc}") from exc

    _render_document_page(result, records[0], 0)
    for record_index, record in enumerate(records[1:], start=1):
        page = Document(template)
        _render_document_page(page, record, record_index)
        break_paragraph = result.add_paragraph()
        break_paragraph.add_run().add_break(WD_BREAK.PAGE)
        body = result._body._element
        for block in _body_blocks(page):
            _insert_before_section_properties(body, block)
    output.parent.mkdir(parents=True, exist_ok=True)
    try:
        result.save(output)
    except OSError as exc:
        raise DiplomaRenderError(f"Could not write rendered DOCX {output}: {exc}") from exc


_ODT_NS = {
    "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
    "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    "style": "urn:oasis:names:tc:opendocument:xmlns:style:1.0",
    "fo": "urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0",
}
for _prefix, _uri in _ODT_NS.items():
    ET.register_namespace(_prefix, _uri)


def _odt_text_nodes(block: ET.Element) -> list[ET.Element]:
    return [node for node in block.iter() if node.text]


def _render_odt_content(content: bytes, record: Mapping[str, Any], record_index: int) -> bytes:
    try:
        root = ET.fromstring(content)
    except ET.ParseError as exc:
        raise DiplomaRenderError(f"Could not read ODT content.xml: {exc}") from exc

    text_root = root.find(".//office:body/office:text", _ODT_NS)
    if text_root is None:
        raise DiplomaRenderError("ODT template does not contain an office text body")
    for block in list(text_root):
        if block.tag not in {f"{{{_ODT_NS['text']}}}p", f"{{{_ODT_NS['text']}}}h"}:
            continue
        nodes = _odt_text_nodes(block)
        combined = "".join(node.text or "" for node in nodes)
        if "{{" not in combined and "}}" not in combined:
            continue
        rendered = _render_text(combined, record, record_index)
        nodes[0].text = rendered
        for node in nodes[1:]:
            node.text = ""
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _odt_page_break_style(root: ET.Element, parent_style: str | None) -> str:
    automatic_styles = root.find(".//office:automatic-styles", _ODT_NS)
    if automatic_styles is None:
        raise DiplomaRenderError("ODT template does not contain automatic styles")
    suffix = re.sub(r"[^A-Za-z0-9_]", "_", parent_style or "Default")
    style_name = f"DiplomaPageBreak_{suffix}"
    style_tag = f"{{{_ODT_NS['style']}}}style"
    if any(item.get(f"{{{_ODT_NS['style']}}}name") == style_name for item in automatic_styles):
        return style_name
    style = ET.Element(style_tag, {
        f"{{{_ODT_NS['style']}}}name": style_name,
        f"{{{_ODT_NS['style']}}}family": "paragraph",
    })
    if parent_style:
        style.set(f"{{{_ODT_NS['style']}}}parent-style-name", parent_style)
    source_style = next((item for item in automatic_styles
                         if item.get(f"{{{_ODT_NS['style']}}}name") == parent_style), None)
    source_properties = None if source_style is None else source_style.find(
        "style:paragraph-properties", _ODT_NS
    )
    if source_properties is not None:
        properties = deepcopy(source_properties)
        properties.set(f"{{{_ODT_NS['fo']}}}break-before", "page")
        style.append(properties)
    else:
        ET.SubElement(style, f"{{{_ODT_NS['style']}}}paragraph-properties", {
            f"{{{_ODT_NS['fo']}}}break-before": "page",
        })
    automatic_styles.append(style)
    return style_name


def _render_odt_records(records: list[Mapping[str, Any]], template_path: str | Path, output_path: str | Path) -> None:
    template = Path(template_path)
    output = Path(output_path)
    if not template.is_file():
        raise DiplomaRenderError(f"Template does not exist: {template}")
    try:
        with ZipFile(template) as source:
            entries = {info.filename: source.read(info.filename) for info in source.infolist()}
    except (OSError, KeyError, ValueError) as exc:
        raise DiplomaRenderError(f"Could not open ODT template {template}: {exc}") from exc
    if "content.xml" not in entries:
        raise DiplomaRenderError(f"ODT template does not contain content.xml: {template}")

    rendered_pages: list[bytes] = []
    for index, record in enumerate(records):
        rendered_pages.append(_render_odt_content(entries["content.xml"], record, index))
    try:
        first_root = ET.fromstring(rendered_pages[0])
        first_text = first_root.find(".//office:body/office:text", _ODT_NS)
        if first_text is None:
            raise DiplomaRenderError("ODT template does not contain an office text body")
        style_attr = f"{{{_ODT_NS['text']}}}style-name"
        for page_xml in rendered_pages[1:]:
            page_root = ET.fromstring(page_xml)
            page_text = page_root.find(".//office:body/office:text", _ODT_NS)
            if page_text is None:
                raise DiplomaRenderError("ODT template does not contain an office text body")
            page_blocks = list(page_text)
            first_block = next((block for block in page_blocks if block.tag in {
                f"{{{_ODT_NS['text']}}}p", f"{{{_ODT_NS['text']}}}h"
            }), None)
            if first_block is not None:
                page_break_style = _odt_page_break_style(first_root, first_block.get(style_attr))
                first_block.set(style_attr, page_break_style)
            for block in page_blocks:
                first_text.append(deepcopy(block))
        entries["content.xml"] = ET.tostring(first_root, encoding="utf-8", xml_declaration=True)
    except ET.ParseError as exc:
        raise DiplomaRenderError(f"Could not assemble ODT output: {exc}") from exc

    output.parent.mkdir(parents=True, exist_ok=True)
    try:
        with ZipFile(output, "w", ZIP_DEFLATED) as destination:
            mimetype = entries.pop("mimetype", None)
            if mimetype is not None:
                destination.writestr("mimetype", mimetype, compress_type=0)
            for name, data in entries.items():
                destination.writestr(name, data)
    except OSError as exc:
        raise DiplomaRenderError(f"Could not write rendered ODT {output}: {exc}") from exc


def _convert_document(source: Path, output: Path, target_format: str) -> None:
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    if converter is None:
        raise DiplomaRenderError("PDF output requires LibreOffice (soffice) to be installed and available on PATH")
    with tempfile.TemporaryDirectory() as temporary:
        temporary_path = Path(temporary)
        try:
            completed = subprocess.run(
                [converter, "--headless", "--convert-to", target_format, "--outdir", str(temporary_path), str(source)],
                capture_output=True,
                text=True,
                check=False,
            )
        except OSError as exc:
            raise DiplomaRenderError(f"Could not run LibreOffice for PDF output: {exc}") from exc
        generated = temporary_path / f"{source.stem}.{target_format}"
        if completed.returncode != 0 or not generated.is_file():
            details = (completed.stderr or completed.stdout).strip()
            raise DiplomaRenderError(f"LibreOffice could not create {target_format.upper()}{': ' + details if details else ''}")
        output.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(generated, output)


def _render_records_to_output(records: list[Mapping[str, Any]], template_path: str | Path, output_path: str | Path) -> None:
    template = Path(template_path)
    output = Path(output_path)
    template_format = template.suffix.lower().lstrip(".")
    output_format = output.suffix.lower().lstrip(".")
    if template_format not in {"docx", "odt"} or output_format not in {"docx", "odt", "pdf"}:
        raise DiplomaRenderError(f"Unsupported template format: {template.suffix or '<none>'}; use .docx or .odt")
    if template_format == "odt":
        native_renderer = _render_odt_records
    else:
        native_renderer = _render_records
    if output_format == template_format:
        native_renderer(records, template, output)
        return
    with tempfile.TemporaryDirectory() as temporary:
        intermediate = Path(temporary) / f"rendered.{template_format}"
        native_renderer(records, template, intermediate)
        _convert_document(intermediate, output, output_format)


def render_diplomas(diplomas_path: str | Path, template_path: str | Path, output_path: str | Path) -> None:
    """Render diploma data from a YAML file into one document."""
    _render_records_to_output(_load_records(Path(diplomas_path)), template_path, output_path)


def render_diploma_data(data: Mapping[str, Any], template_path: str | Path, output_path: str | Path) -> None:
    """Render already-generated diploma data without an intermediate YAML file."""
    _render_records_to_output(_records_from_data(data), template_path, output_path)
