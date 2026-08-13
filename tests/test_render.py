from pathlib import Path
import sys
import zipfile

import yaml
from docx import Document

from practiscore_diplomas.render import DiplomaRenderError, render_diplomas
from practiscore_diplomas.cli import main
from tests.test_cli import write_config, write_export


def write_template(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("{{ first_")
    paragraph.add_run("line }} / {{ shooter.last_name }}")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{ place }} {{ shooter.raw_time }} {{ third_line }} {{ fourth_line }}"
    document.save(path)


def write_data(path: Path) -> None:
    path.write_text(
        yaml.safe_dump(
            {
                "diplomas": {
                    "first_series": [
                        {
                            "place": 1,
                            "first_line": "FIRST",
                            "second_line": "Alpha",
                            "shooter": {"last_name": "One", "raw_time": 12.5},
                        },
                        {
                            "place": 2,
                            "first_line": "SECOND",
                            "second_line": "Beta",
                            "third_line": "EXTRA",
                            "shooter": {"last_name": "Two", "raw_time": 15},
                        },
                    ],
                    "second_series": [
                        {
                            "place": 1,
                            "first_line": "THIRD",
                            "second_line": "Gamma",
                            "shooter": {"last_name": "Three", "raw_time": 18},
                        }
                    ],
                }
            },
            sort_keys=False,
        ),
        encoding="utf-8",
    )


def write_odt_template(path: Path) -> None:
    content = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0">
  <office:automatic-styles><style:style style:name="Centered" style:family="paragraph"><style:paragraph-properties fo:text-align="center"/></style:style></office:automatic-styles>
  <office:body><office:text><text:p text:style-name="Centered">{{ first_line }} / {{ shooter.last_name }}</text:p><text:p>{{ place }}</text:p></office:text></office:body>
</office:document-content>"""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("mimetype", "application/vnd.oasis.opendocument.text")
        archive.writestr("content.xml", content)


def test_render_creates_one_page_per_diploma_and_preserves_order(tmp_path: Path):
    template = tmp_path / "template.docx"
    data = tmp_path / "diplomas.yaml"
    output = tmp_path / "result.docx"
    write_template(template)
    write_data(data)

    render_diplomas(data, template, output)

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "FIRST / One" in text
    assert "SECOND / Two" in text
    assert "THIRD / Three" in text
    assert text.index("FIRST") < text.index("SECOND") < text.index("THIRD")
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "1 12.5  " in table_text
    assert "2 15 EXTRA " in table_text
    assert "1 18  " in table_text

    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert document_xml.count('w:type="page"') == 2

    rendered_paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert rendered_paragraphs == ["FIRST / One", "SECOND / Two", "THIRD / Three"]


def test_render_rejects_unknown_placeholder(tmp_path: Path):
    template = tmp_path / "template.docx"
    data = tmp_path / "diplomas.yaml"
    output = tmp_path / "result.docx"
    document = Document()
    document.add_paragraph("{{ missing_field }}")
    document.save(template)
    write_data(data)

    try:
        render_diplomas(data, template, output)
    except DiplomaRenderError as exc:
        assert "unknown placeholder" in str(exc)
    else:
        raise AssertionError("render_diplomas should reject unknown placeholders")


def test_render_supports_odt_template(tmp_path: Path):
    template = tmp_path / "template.odt"
    data = tmp_path / "diplomas.yaml"
    output = tmp_path / "result.odt"
    write_odt_template(template)
    write_data(data)

    render_diplomas(data, template, output)

    with zipfile.ZipFile(output) as archive:
        content = archive.read("content.xml").decode("utf-8")
    assert "FIRST / One" in content
    assert "SECOND / Two" in content
    assert "THIRD / Three" in content
    assert "DiplomaPageBreak" in content
    assert "parent-style-name=\"Centered\"" in content
    assert "fo:text-align=\"center\"" in content
    assert "{{" not in content


def test_render_cli_accepts_output_odt(tmp_path: Path, monkeypatch):
    template = tmp_path / "template.odt"
    data = tmp_path / "diplomas-data.yaml"
    output = tmp_path / "diplomas.odt"
    write_odt_template(template)
    write_data(data)
    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(sys, "argv", ["practiscore-diplomas", "render", "-d", str(data), "-t", str(template), "--output-odt", str(output)])

    assert main() == 0
    assert output.is_file()


def test_render_odt_template_defaults_to_odt_output(tmp_path: Path, monkeypatch):
    template = tmp_path / "template.odt"
    data = tmp_path / "diplomas-data.yaml"
    write_odt_template(template)
    write_data(data)
    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(sys, "argv", ["practiscore-diplomas", "render", "-d", str(data), "-t", str(template)])

    assert main() == 0
    assert (tmp_path / "diplomas.odt").is_file()


def test_render_output_option_can_use_default_path(tmp_path: Path, monkeypatch):
    template = tmp_path / "template.odt"
    data = tmp_path / "diplomas-data.yaml"
    write_odt_template(template)
    write_data(data)
    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(sys, "argv", ["practiscore-diplomas", "render", "-d", str(data), "-t", str(template), "--output-odt"])

    assert main() == 0
    assert (tmp_path / "diplomas.odt").is_file()


def test_render_cli_requires_diplomas_data_input(tmp_path: Path, monkeypatch):
    template = tmp_path / "template.docx"
    data = tmp_path / "diplomas-data.yaml"
    output = tmp_path / "diplomas.docx"
    write_template(template)
    write_data(data)
    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(sys, "argv", ["practiscore-diplomas", "render", "-d", str(data), "-t", str(template), "-o", str(output)])

    assert main() == 0
    assert output.is_file()


def test_render_cli_can_parse_match_and_uses_match_name_output(tmp_path: Path, monkeypatch):
    input_dir = tmp_path / "match"
    template = tmp_path / "template.docx"
    config = tmp_path / "config.yaml"
    write_export(input_dir)
    write_config(config)
    write_template(template)
    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(sys, "argv", ["practiscore-diplomas", "render", "-i", str(input_dir), "-c", str(config), "-t", str(template)])

    assert main() == 0
    assert (tmp_path / "diplomas-Synthetic-Match.docx").is_file()
    assert (tmp_path / "shooters-summary-Synthetic-Match.yaml").is_file()
    assert (tmp_path / "diplomas-data-Synthetic-Match.yaml").is_file()
